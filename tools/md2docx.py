"""Convert the MVP 3 markdown spec to .docx.

Deliberately handles only the constructs that document actually uses -- headings,
paragraphs, bold, inline code, fenced code blocks, pipe tables, bullet and numbered
lists, blockquotes and horizontal rules. A general markdown converter would be more
code and more ways to be subtly wrong about a document nobody would re-check.
"""
import re
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

SRC, OUT = sys.argv[1], sys.argv[2]

doc = Document()

# Base style: a document that will be read on screen and printed, so 10.5pt Calibri
# rather than Word's default 11pt Cambria.
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.08

CODE_FONT = "Consolas"
CODE_BG = "F2F2F2"
RULE_GREY = RGBColor(0x80, 0x80, 0x80)


def shade(target, hexfill):
    """Apply a background fill. Word has no paragraph 'background', only shading.

    python-docx exposes wrappers, not the XML: a cell's element is `_tc` and a
    paragraph's is `_p`, and neither wrapper carries the get_or_add_* methods.
    """
    if hasattr(target, "_tc"):
        pr = target._tc.get_or_add_tcPr()
    else:
        pr = target._p.get_or_add_pPr()
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:color"), "auto")
    el.set(qn("w:fill"), hexfill)
    pr.append(el)


# Inline markup: **bold**, `code`, [text](target). Split on all three at once so a
# bold span containing code still renders both -- handling them in sequence would
# make the second pass see already-consumed text.
INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")


def add_inline(par, text):
    for piece in INLINE.split(text):
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**"):
            par.add_run(piece[2:-2]).bold = True
        elif piece.startswith("`") and piece.endswith("`"):
            r = par.add_run(piece[1:-1])
            r.font.name = CODE_FONT
            r.font.size = Pt(9.5)
        elif piece.startswith("[") and "](" in piece:
            label = piece[1 : piece.index("]")]
            target = piece[piece.index("](") + 2 : -1]
            # Rendered as "label (target)" rather than a live hyperlink: the targets are
            # repo-relative paths and issue numbers, which resolve in GitHub and not in
            # Word, so a clickable link would be a broken promise.
            r = par.add_run(label)
            r.font.name = CODE_FONT
            r.font.size = Pt(9.5)
            if target != label:
                par.add_run(" (%s)" % target).italic = True
        else:
            par.add_run(piece)


def add_code_block(lines):
    par = doc.add_paragraph()
    par.paragraph_format.left_indent = Pt(14)
    par.paragraph_format.space_before = Pt(4)
    par.paragraph_format.space_after = Pt(8)
    shade(par, CODE_BG)
    run = par.add_run("\n".join(lines))
    run.font.name = CODE_FONT
    run.font.size = Pt(9)


def add_table(rows):
    header, body = rows[0], rows[1:]
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, cell_text in enumerate(header):
        c = t.rows[0].cells[i]
        c.text = ""
        shade(c, "E7E6E6")
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        add_inline(p, cell_text)
        for r in p.runs:
            r.bold = True
    for row in body:
        cells = t.add_row().cells
        for i, cell_text in enumerate(row[: len(header)]):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            add_inline(p, cell_text)
    doc.add_paragraph()


def split_row(line):
    # Strip the outer pipes, then split on unescaped pipes. The spec uses \| inside
    # cells for literal pipes in type unions like "object \| null".
    inner = line.strip().strip("|")
    parts = re.split(r"(?<!\\)\|", inner)
    return [p.strip().replace("\\|", "|") for p in parts]


src = open(SRC, encoding="utf-8").read().replace("\r\n", "\n").split("\n")

i = 0
pending_table = []


def flush_table():
    global pending_table
    if pending_table:
        add_table(pending_table)
        pending_table = []


while i < len(src):
    line = src[i]

    # fenced code
    if line.startswith("```"):
        buf = []
        i += 1
        while i < len(src) and not src[i].startswith("```"):
            buf.append(src[i])
            i += 1
        i += 1
        flush_table()
        add_code_block(buf)
        continue

    # table rows accumulate; the |---|---| separator is skipped
    if line.strip().startswith("|") and line.strip().endswith("|"):
        cells = split_row(line)
        if not all(re.fullmatch(r":?-{2,}:?", c) for c in cells if c):
            pending_table.append(cells)
        i += 1
        continue
    flush_table()

    stripped = line.strip()

    if not stripped:
        i += 1
        continue

    if stripped in ("---", "***", "___"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("• • •")
        r.font.color.rgb = RULE_GREY
        i += 1
        continue

    m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
    if m:
        level = len(m.group(1))
        text = m.group(2)
        if level == 1:
            h = doc.add_heading(level=0)
            add_inline(h, text)
        else:
            h = doc.add_heading(level=min(level - 1, 4))
            add_inline(h, text)
        i += 1
        continue

    if stripped.startswith("> "):
        buf = [stripped[2:]]
        i += 1
        while i < len(src) and src[i].strip().startswith("> "):
            buf.append(src[i].strip()[2:])
            i += 1
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.space_before = Pt(4)
        add_inline(p, " ".join(buf))
        for r in p.runs:
            r.italic = True
        continue

    m = re.match(r"^(\s*)[-*]\s+(.*)$", line)
    if m:
        depth = len(m.group(1)) // 2
        p = doc.add_paragraph(style="List Bullet" if depth == 0 else "List Bullet 2")
        add_inline(p, m.group(2))
        i += 1
        continue

    m = re.match(r"^(\s*)\d+\.\s+(.*)$", line)
    if m:
        p = doc.add_paragraph(style="List Number")
        add_inline(p, m.group(2))
        i += 1
        continue

    # plain paragraph: join wrapped lines until a blank or a block starter
    buf = [stripped]
    i += 1
    while i < len(src):
        nxt = src[i]
        if not nxt.strip():
            break
        if re.match(r"^(#{1,6}\s|\s*[-*]\s|\s*\d+\.\s|>\s|```)", nxt) or nxt.strip().startswith("|"):
            break
        if nxt.strip() in ("---", "***", "___"):
            break
        buf.append(nxt.strip())
        i += 1
    add_inline(doc.add_paragraph(), " ".join(buf))

flush_table()
doc.save(OUT)
print("wrote", OUT)
